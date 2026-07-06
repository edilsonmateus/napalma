from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "MANUAL_MESTRE_77GIRA.md"
OUTPUT = ROOT / "Manual_Mestre_77Gira.docx"

NAVY = "18243A"
BLUE = "276EF1"
CYAN = "23A6D5"
GOLD = "EAAA00"
INK = "202631"
MUTED = "667085"
LIGHT = "E8EEF5"
PALE = "F5F7FA"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    total = sum(widths_dxa)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def font(run, name="Aptos", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_inline(paragraph, text, size=10.5, color=INK, bold=False):
    text = re.sub(r"\[([^]]+)\]\([^)]+\)", r"\1", text)
    tokens = re.split(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for token in tokens:
        if not token:
            continue
        is_code = token.startswith("`") and token.endswith("`")
        is_bold = token.startswith("**") and token.endswith("**")
        is_italic = token.startswith("*") and token.endswith("*") and not is_bold
        clean = token[1:-1] if is_code or is_italic else token[2:-2] if is_bold else token
        run = paragraph.add_run(clean)
        font(run, "Consolas" if is_code else "Aptos", 9.5 if is_code else size, color, bold or is_bold, is_italic)
        if is_code:
            run.font.highlight_color = None


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color, before, after in (
        ("Heading 1", 17, NAVY, 18, 9),
        ("Heading 2", 13.5, BLUE, 14, 7),
        ("Heading 3", 11.5, NAVY, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Aptos"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.18

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    code.font.name = "Consolas"
    code.font.size = Pt(8.5)
    code.font.color.rgb = RGBColor.from_string(NAVY)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(3)


def set_page_furniture(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.88)
    section.right_margin = Inches(0.88)
    section.header_distance = Inches(0.34)
    section.footer_distance = Inches(0.34)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("77GIRA  /  MANUAL MESTRE")
    font(r, "Aptos", 8, MUTED, True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("77Gira  •  Produto, operação e publicidade   |   ")
    font(r, "Aptos", 8, MUTED)
    add_field(p, "PAGE")


def add_cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("77GIRA")
    font(r, "Aptos Display", 13, GOLD, True)
    p.paragraph_format.space_after = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Manual Mestre")
    font(r, "Aptos Display", 32, NAVY, True)
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Produto, agentes, jornadas, operação e publicidade")
    font(r, "Aptos", 15, BLUE)
    p.paragraph_format.space_after = Pt(30)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Do acesso público sem login à gestão de artistas, casas, produtores e campanhas de marca.")
    font(r, "Aptos", 11, MUTED, italic=True)
    p.paragraph_format.left_indent = Inches(0.65)
    p.paragraph_format.right_indent = Inches(0.65)
    p.paragraph_format.space_after = Pt(80)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("VERSÃO 1.0  •  5 DE JULHO DE 2026")
    font(r, "Aptos", 9, MUTED, True)
    doc.add_page_break()


def add_table(doc, rows):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    widths = [9360 // cols] * cols
    widths[-1] += 9360 - sum(widths)
    set_table_geometry(table, widths)
    for i, row in enumerate(rows):
        for j in range(cols):
            cell = table.cell(i, j)
            value = row[j] if j < len(row) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_inline(p, value, 8.5 if cols >= 4 else 9, WHITE if i == 0 else INK, i == 0)
            if i == 0:
                set_cell_shading(cell, NAVY)
            elif i % 2 == 0:
                set_cell_shading(cell, PALE)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def parse_markdown(doc, text):
    lines = text.splitlines()
    # The Word cover replaces the Markdown title block and metadata.
    start = next(i for i, line in enumerate(lines) if line.startswith("## 1. "))
    lines = lines[start:]
    i = 0
    in_code = False
    while i < len(lines):
        line = lines[i].rstrip()
        if line.startswith("```"):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            p = doc.add_paragraph(style="Code Block")
            add_inline(p, line, 8.5, NAVY)
            i += 1
            continue
        if line == "---":
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|?\s*:?-+", lines[i + 1]):
            rows = []
            rows.append([c.strip() for c in line.strip("|").split("|")])
            i += 2
            while i < len(lines) and lines[i].startswith("|"):
                rows.append([c.strip() for c in lines[i].strip("|").split("|")])
                i += 1
            add_table(doc, rows)
            continue
        heading = re.match(r"^(#{2,4})\s+(.+)$", line)
        if heading:
            level = min(len(heading.group(1)) - 1, 3)
            title = heading.group(2)
            p = doc.add_paragraph(style=f"Heading {level}")
            add_inline(p, title, 17 if level == 1 else 13.5 if level == 2 else 11.5, NAVY if level != 2 else BLUE, True)
            i += 1
            continue
        bullet = re.match(r"^\s*[-*]\s+(.+)$", line)
        number = re.match(r"^\s*\d+[.)]\s+(.+)$", line)
        if bullet or number:
            p = doc.add_paragraph(style="List Bullet" if bullet else "List Number")
            add_inline(p, (bullet or number).group(1))
            i += 1
            continue
        if line.strip():
            p = doc.add_paragraph()
            add_inline(p, line.strip())
        i += 1


def set_document_properties(doc):
    props = doc.core_properties
    props.title = "Manual Mestre do 77Gira"
    props.subject = "Produto, agentes, jornadas, operação e publicidade"
    props.author = "77 GiraMundo"
    props.keywords = "77Gira, manual, artistas, casas, produtores, publicidade, ADS"
    props.comments = "Versão consolidada em 5 de julho de 2026"


def main():
    source = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    configure_styles(doc)
    set_page_furniture(doc)
    set_document_properties(doc)
    add_cover(doc)
    parse_markdown(doc, source)
    doc.settings.update_fields_on_open = True
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
